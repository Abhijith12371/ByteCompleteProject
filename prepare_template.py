import os
import docx

def main():
    try:
        doc = docx.Document('template.docx')
    except Exception as e:
        print(f"Failed to open template: {e}")
        return

    # 1. Replace "Rama Krishna" and "7 B" in the first table
    info_table = doc.tables[0]
    info_table.cell(0, 1).text = "{{ student.name }}"
    info_table.cell(1, 1).text = "{{ student.class_name }}"

    # 2. Observations Table (table 1)
    obs_table = doc.tables[1]
    row1 = obs_table.rows[1]
    row1.cells[0].text = "{% tr for obs in student.filtered_obs %}{{ obs.activity }}"
    row1.cells[1].text = "{{ obs.engagement }}%"
    row1.cells[2].text = "{{ obs.interest_level }}"
    row1.cells[3].text = "{{ obs.observation }}{% endtr %}"

    # Delete other rows in observations table
    for row in obs_table.rows[2:]:
        obs_table._element.remove(row._tr)

    # 3. Recommendations Table (table 2)
    rec_table = doc.tables[2]
    row1 = rec_table.rows[1]
    row1.cells[0].text = "{% tr for rec in student.recs %}{{ rec.area }} ({{ rec.engagement }}%)"
    row1.cells[1].text = "{{ rec.future_pathways }}"
    row1.cells[2].text = "{{ rec.support_activities }}{% endtr %}"

    for row in rec_table.rows[2:]:
        rec_table._element.remove(row._tr)

    # 4. Tips
    tips_start = -1
    for i, p in enumerate(doc.paragraphs):
        if "How Parents Can Help" in p.text:
            tips_start = i
            break

    if tips_start != -1 and tips_start + 1 < len(doc.paragraphs):
        p1 = doc.paragraphs[tips_start + 1]
        p1.text = "{% p for tip in student.parent_tips %}{{ loop.index }}. {{ tip }}{% endp %}"
        
        # Remove following tips
        body = doc.element.body
        try:
            body.remove(doc.paragraphs[tips_start + 2]._element)
            body.remove(doc.paragraphs[tips_start + 3]._element)
        except:
            pass

    # 5. Conclusion
    for p in doc.paragraphs:
        if p.text.startswith("Your child’s strengths") or p.text.startswith("Conclusion"):
            # Wait, the conclusion text is in the paragraph after "Conclusion:"
            pass
            
    # Let's find Conclusion:
    conc_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Conclusion:" in p.text:
            conc_idx = i
            break
            
    if conc_idx != -1 and conc_idx + 1 < len(doc.paragraphs):
        doc.paragraphs[conc_idx + 1].text = "{{ student.conclusion }}"

    # 6. Replace Spider Chart Image Placeholder
    # We will insert it before the Observations heading
    for p in doc.paragraphs:
        if "Observations:" in p.text:
            p.insert_paragraph_before("{{ student.spider_chart }}")
            break

    # 7. Remove existing static spider chart.
    for p in doc.paragraphs:
        for r in p.runs:
            if 'graphic' in r._element.xml or 'pic' in r._element.xml:
                p._element.remove(r._element)

    # 8. Wrap the entire thing in a loop for multiple students
    # But wait, it's easier to use a single page break between students.
    # The docxtpl supports standard Jinja {% for student in students %}...{% endfor %}
    # We insert {% p for student in students %} at the very top
    p_top = doc.paragraphs[0]
    p_top.insert_paragraph_before("{% p for student in students %}")
    
    # Add {% p if not loop.last %}{{ page_break }}{% endif %} at the end
    last_p1 = doc.add_paragraph("{% p if not loop.last %}{{ page_break }}{% endif %}")
    last_p2 = doc.add_paragraph("{% p endfor %}")

    doc.save('template_jinja.docx')
    print("template_jinja.docx created successfully.")

if __name__ == '__main__':
    main()
